from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _safe_list(x: Any) -> List[str]:
    if isinstance(x, list):
        return [str(i) for i in x if str(i).strip()]
    return []


def _safe_str(x: Any, default: str = "не указано") -> str:
    s = str(x).strip() if x is not None else ""
    return s if s else default


def _make_bold(cell) -> None:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True


def build_protocol_docx(protocol: Dict[str, Any], output_path: Path) -> None:
    """
    Build a Word (.docx) meeting protocol from structured data.

    Expected protocol structure:
    {
      "date": str,
      "time": str,
      "format": str,
      "topic": str,
      "participants": [str],
      "agenda": [str],
      "key_points": [str],
      "decisions": [{"decision": str, "context": str}],
      "action_items": [{"task": str, "owner": str, "due": str, "done_criteria": str}],
      "risks": [{"risk": str, "mitigation": str}],
      "open_questions": [str],
      "next_meeting": {"datetime": str, "topic": str}
    }
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Base font for a clean corporate look
    try:
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
    except Exception:
        pass

    # Title
    title = doc.add_heading("ПРОТОКОЛ СОВЕЩАНИЯ", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. General info
    doc.add_heading("1. Общая информация", level=2)
    doc.add_paragraph(f"Дата: {_safe_str(protocol.get('date'))}")
    doc.add_paragraph(f"Время: {_safe_str(protocol.get('time'))}")
    doc.add_paragraph(f"Формат: {_safe_str(protocol.get('format'))}")
    doc.add_paragraph(f"Тема: {_safe_str(protocol.get('topic'))}")

    doc.add_paragraph("Участники:")
    participants = _safe_list(protocol.get("participants")) or ["не указано"]
    for p in participants:
        doc.add_paragraph(p, style="List Bullet")

    # 2. Agenda
    doc.add_heading("2. Повестка (что обсуждали)", level=2)
    agenda = _safe_list(protocol.get("agenda"))
    if agenda:
        for item in agenda:
            doc.add_paragraph(item, style="List Number")
    else:
        doc.add_paragraph("не указано")

    # 3. Key points
    doc.add_heading("3. Ключевые тезисы обсуждения (кратко)", level=2)
    key_points = _safe_list(protocol.get("key_points"))
    if key_points:
        for item in key_points:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("не указано")

    # 4. Decisions
    doc.add_heading("4. Принятые решения", level=2)
    decisions = protocol.get("decisions") if isinstance(protocol.get("decisions"), list) else []
    if decisions:
        for d in decisions:
            if not isinstance(d, dict):
                continue
            decision = _safe_str(d.get("decision"))
            context = _safe_str(d.get("context"))
            doc.add_paragraph(f"Решение: {decision}")
            doc.add_paragraph(f"Основание/контекст: {context}")
            doc.add_paragraph("")  # spacer
    else:
        doc.add_paragraph("Решения по итогам встречи не зафиксированы.")

    # 5. Action items table
    doc.add_heading("5. Задачи и поручения (Action Items)", level=2)
    action_items = protocol.get("action_items") if isinstance(protocol.get("action_items"), list) else []

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.autofit = True

    hdr = table.rows[0].cells
    hdr[0].text = "№"
    hdr[1].text = "Задача"
    hdr[2].text = "Ответственный"
    hdr[3].text = "Срок"
    hdr[4].text = "Критерий готовности / результат"
    for c in hdr:
        _make_bold(c)

    if action_items:
        for i, a in enumerate(action_items, start=1):
            if not isinstance(a, dict):
                continue
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = _safe_str(a.get("task"))
            row[2].text = _safe_str(a.get("owner"), "не указан")
            row[3].text = _safe_str(a.get("due"), "не указан")
            row[4].text = _safe_str(a.get("done_criteria"))
    else:
        row = table.add_row().cells
        row[0].text = "1"
        row[1].text = "Задачи не зафиксированы"
        row[2].text = "—"
        row[3].text = "—"
        row[4].text = "—"

    # 6. Risks
    doc.add_heading("6. Риски и блокеры", level=2)
    risks = protocol.get("risks") if isinstance(protocol.get("risks"), list) else []
    if risks:
        for r in risks:
            if not isinstance(r, dict):
                continue
            risk = _safe_str(r.get("risk"))
            mitigation = _safe_str(r.get("mitigation"))
            doc.add_paragraph(f"{risk} — {mitigation}", style="List Bullet")
    else:
        doc.add_paragraph("Риски и блокеры не зафиксированы.")

    # 7. Open questions
    doc.add_heading("7. Открытые вопросы", level=2)
    oq = _safe_list(protocol.get("open_questions"))
    if oq:
        for q in oq:
            doc.add_paragraph(q, style="List Bullet")
    else:
        doc.add_paragraph("Открытые вопросы не зафиксированы.")

    # 8. Next meeting
    doc.add_heading("8. Следующая встреча", level=2)
    nm = protocol.get("next_meeting") if isinstance(protocol.get("next_meeting"), dict) else {}
    doc.add_paragraph(f"Дата/время: {_safe_str(nm.get('datetime'))}")
    doc.add_paragraph(f"Предварительная тема: {_safe_str(nm.get('topic'))}")

    doc.save(output_path)
